import docx2txt
import pandas as pd
import docx
from docx import Document
import nltk as nl
import pandas as pd
import numpy as np
my_text = docx2txt.process("sample1_cv.docx")
document = Document('sample1_cv.docx')
lst = pd.read_csv("repository.csv")
a=my_text.encode("ISO-8859-1","ignore")
headings=[]
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                headings.append(paragraph.text)
headings=[str(k) for k in headings]
sections=a.split('\n\n')
sec_no=0
paras=['']
for sec in sections:
    if sec in headings:
        sec_no+=1
        paras.append('')
        continue
    paras[sec_no]+=sec+"\n"
import nltk
tokens=[]
token_final=[]
for i,p in enumerate(paras):
    #print p
    tokens.append(nltk.tokenize.word_tokenize(p))
#xy =nltk.tokenize.word_tokenize(paras)
for token in tokens:
    token_final.append(token)
paras_sent=[]
for para in paras:
    paras_sent.append([])
    paras_sent[-1].extend(nl.tokenize.sent_tokenize(para.replace('\n','.')))
paras_sent=[]
for para in paras:
    paras_sent.append([])
    paras_sent[-1].extend(para.split('\n'))
repo=[str(list(k)[0]).lower() for k in list(np.array(lst))]
from difflib import SequenceMatcher
#from  sets import Set
def similar(a,b):
    return SequenceMatcher(None,a,b).ratio()

repo=[str(list(k)[0]).lower() for k in list(np.array(lst))]

personal=[]
for token in tokens:
    for retoken in token:
        for r in repo:
            x=retoken.lower()
            # print retoken
            if similar(x,r) > 0.9:
                if r not in personal:
                    personal.append(r)
a=a.lower()
for i in range(0,len(repo)):
    st=str(repo[i]).lower()
    if a.find(st)>=0:
        if repo[i] not in personal:
            personal.append(st)
                    
Technical_skill=personal
print "The technical skills are"
for i in range(0,len(Technical_skill)):
    print "{0}.{1}".format(i+1,Technical_skill[i])
gen_heading = ['project','research','internship']
under_heads=[]
for h in gen_heading:
    for i,hed in enumerate(headings):
        x=hed.lower()
        if similar(h,x) > 0.5:
            #print x
            under_heads.append(i)
            

to_search=[]
for heads in under_heads:
    #print heads
    start=a.index(headings[heads].lower())
    end=-1
    if heads+1 < len(headings):
        end=a.index(headings[heads+1].lower())
    to_search.append(a[start:end])
    
keys=[]
for j,sec in enumerate(to_search):
    for i in range(len(personal)):
            st=str(sec).lower()
            if st.find(" "+personal[i]+" ")>=0 or st.find("."+personal[i]+" ")>=0:
                    #print j,personal[i]
                    keys.append(personal[i])

skills_read=[]
occurence={}
for skil in keys:
    if skil in skills_read:
        continue
    num=keys.count(skil)
    if num in occurence.keys():
        occurence[num].append(skil)
    else:
        occurence[num]=[]
        occurence[num].append(skil)
    skills_read.append(skil)
for skil in personal:
    if skil in skills_read:
        continue
    num=keys.count(skil)
    if num in occurence.keys():
        occurence[num].append(skil)
    else:
        occurence[num]=[]
        occurence[num].append(skil)
    skills_read.append(skil)

occur=sorted(occurence.items())
occur.reverse()
skills_ranked=[]
for occ in occur:
    skills_ranked.extend(occ[1])
print "\n\nThe skill ranking is as follows"
for i in range(0,len(skills_ranked)):
    print "{0}.{1}".format(i+1,skills_ranked[i])
